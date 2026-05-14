import os
import re
import io
import json
import requests
import PyPDF2
import docx
from openai import OpenAI
from flask import current_app

SYSTEM_PROMPT = """
Ты специалист по составлению регламентов. Работаешь в компании RedCat.
Тебе нужно составлять регламенты фиксации и бронирования клиентов на основе предоставленного документа (агентского договора или приложений к нему).

ФОРМА (строго JSON):
{
  "internal_regulation": {
    "company": "",
    "fixation_address": "",
    "fixation_period": "",
    "commerce_fixation": "",
    "foreign_numbers_fixation": "",
    "cross_fixation": "",
    "fixation_start": "",
    "uniqueness_extension": "",
    "viewing_appointment": "",
    "escort_required": "",
    "inspection_report_required": "",
    "developer_emails": "",
    "response_time": "",
    "notes": ""
  },
  "booking_regulation": {
    "how_to_update_tariffs": "",
    "how_to_update_remains": ""
  }
}

Твоя задача:
1. Внимательно прочитай текст документа.
2. Для каждого поля формы попытайся найти ЯВНОЕ указание.
   - Если нашёл – запиши точную формулировку (или краткую суть).
3. Если явного указания нет, но есть КОСВЕННЫЕ данные (например, упоминается похожий процесс, но без деталей), напиши:
   "Предположительно: <кратко что именно>. Уточнить: <конкретный вопрос, который нужно задать застройщику>."
4. Если по данному пункту нет НИКАКОЙ информации, напиши:
   "Уточнить: <что именно нужно уточнить (например, адрес фиксации, срок фиксации, контакты)>."
5. НЕ придумывай данные, которых нет в тексте.
6. Отвечай ТОЛЬКО JSON, без каких-либо пояснений до или после.
"""

def extract_text_from_url(url):
    """Извлекает текст из файла, доступного по URL, включая таблицы DOCX."""
    try:
        response = requests.get(url)
        response.raise_for_status()
        filename = url.split('/')[-1]
        ext = os.path.splitext(filename)[1].lower()
        file_content = io.BytesIO(response.content)

        if ext == '.pdf':
            reader = PyPDF2.PdfReader(file_content)
            text = ''
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + '\n'
            if not text.strip():
                raise ValueError("PDF не содержит текстового слоя (возможно, скан).")
            return text

        elif ext in ('.docx', '.doc'):
            doc = docx.Document(file_content)
            full_text = []

            # Обычные параграфы
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text.strip())

            # Таблицы
            for table in doc.tables:
                for row in table.rows:
                    row_data = []
                    for cell in row.cells:
                        row_data.append(cell.text.strip())
                    full_text.append(' | '.join(row_data))

            # Колонтитулы
            for section in doc.sections:
                header = section.header
                if header:
                    for para in header.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text.strip())
                footer = section.footer
                if footer:
                    for para in footer.paragraphs:
                        if para.text.strip():
                            full_text.append(para.text.strip())

            return '\n'.join(full_text)

        elif ext == '.txt':
            return response.text

        else:
            raise ValueError(f"Неподдерживаемый формат файла: {ext}")

    except Exception as e:
        raise ValueError(f"Не удалось извлечь текст из документа: {e}")

def process_agency_agreement(url):
    """Обрабатывает агентский договор, доступный по публичному URL."""
    try:
        text = extract_text_from_url(url)
        if not text.strip():
            raise ValueError("Извлечённый текст пуст.")
    except Exception as e:
        current_app.logger.error(f"[AI AGENT] Ошибка извлечения текста: {e}")
        raise

    # Логируем первые 800 символов
    current_app.logger.info(f"[AI AGENT] Extracted text (first 800 chars): {text[:800]}")

    client = OpenAI(
        api_key=current_app.config['OPENROUTER_API_KEY'],
        base_url="https://openrouter.ai/api/v1",
        default_headers={
            "HTTP-Referer": "https://builder-admin-q7mb.onrender.com",
            "X-Title": "RedCat CRM"
        }
    )

    try:
        response = client.chat.completions.create(
            model="openrouter/free",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": f"Документ:\n{text[:15000]}"}
            ],
            temperature=0.1,
            max_tokens=2000,
        )
        answer = response.choices[0].message.content
        current_app.logger.info(f"[AI AGENT] Raw answer (first 800 chars): {answer[:800]}")
    except Exception as e:
        current_app.logger.error(f"[AI AGENT] Ошибка вызова OpenRouter: {e}")
        raise

    # Пытаемся извлечь JSON
    try:
        json_match = re.search(r'\{.*\}', answer, re.DOTALL)
        if json_match:
            json_str = json_match.group()
            data = json.loads(json_str)
            return data
        else:
            current_app.logger.warning("[AI AGENT] JSON не найден в ответе, возвращаю raw_response.")
            return {"raw_response": answer}
    except Exception as e:
        current_app.logger.error(f"[AI AGENT] Ошибка парсинга JSON: {e}")
        return {"raw_response": answer}
